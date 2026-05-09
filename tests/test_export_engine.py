import json
import os
from io import BytesIO
from pathlib import Path

os.environ.setdefault('WORKING_DIR', str(Path(__file__).resolve().parents[1]))

from fastapi import HTTPException
from docx import Document
from openpyxl import load_workbook

from kuro_backend import chat_history, compliance_db, finance_db, intelligence_db
from kuro_backend.export_engine.export_models import ExportPayload, ExportRequest
from kuro_backend.export_engine.exporters.csv_exporter import CsvExporter
from kuro_backend.export_engine.exporters.docx_exporter import DocxExporter
from kuro_backend.export_engine.export_security import sanitize_export_payload, validate_export_permission
from kuro_backend.export_engine.exporters.json_exporter import JsonExporter
from kuro_backend.export_engine.exporters.markdown_exporter import MarkdownExporter
from kuro_backend.export_engine.exporters.pdf_exporter import PdfExporter
from kuro_backend.export_engine.exporters.txt_exporter import TxtExporter
from kuro_backend.export_engine.exporters.xlsx_exporter import XlsxExporter
from kuro_backend.export_engine.renderers.chat_renderer import render_chat_session, render_selected_messages
from kuro_backend.export_engine.renderers.compliance_renderer import render_compliance_report
from kuro_backend.export_engine.renderers.finance_renderer import render_market_snapshot
from kuro_backend.export_engine.renderers.intelligence_renderer import render_intelligence_report


def _seed_chat(chat_id='chat_export_1', username='Pantronux', persona='advisor'):
    chat_history.create_session(chat_id, username, persona, title='Export Session')
    chat_history.add_message('web', 'user', 'First user message', [], persona, None, username, chat_id)
    chat_history.add_message('web', 'assistant', json.dumps({'summary': 'Structured response'}), ['brief.json'], persona, None, username, chat_id)
    chat_history.add_message('web', 'user', 'Third user message', [], persona, None, username, chat_id)
    return chat_history.get_history(chat_id=chat_id, username=username, limit=9999)


def _seed_intelligence(username='Pantronux', briefing_date='2026-05-10'):
    intelligence_db.save_briefing(
        briefing_date,
        'Executive summary',
        {'status_pagi': 'stable', 'intelijen_sektoral': 'energy up'},
        ['signal-a', 'signal-b'],
        [{'symbol': 'BBCA.JK', 'stance': 'watch', 'score': 8}],
        username=username,
    )
    intelligence_db.save_research_sources(
        session_id='sess-1',
        username=username,
        chat_id='chat-1',
        sources=[{
            'query': 'energy outlook',
            'source_type': 'news',
            'title': 'Energy Outlook',
            'link': 'https://example.com/energy',
            'snippet': 'Energy trend remains positive',
            'year': 2026,
            'cited_by': 0,
        }],
    )


def _seed_compliance():
    compliance_db.add_evidence('evidence_a.pdf', '/tmp/evidence_a.pdf', 'policy', 'iso27001', 'A.5')
    rows = compliance_db.get_evidence_matrix('iso27001')
    compliance_db.update_evidence_status(rows[0]['id'], 'compliant', 'good', 'keep')
    compliance_db.add_audit_trail('review', 'reviewed A.5', 'iso27001')


def _seed_market(username='Pantronux'):
    finance_db.upsert_watched_symbol('BBCA.JK', 'BBCA', username=username)
    finance_db.apply_watched_price('BBCA.JK', 1000, username=username)
    finance_db.apply_watched_price('BBCA.JK', 1025, username=username)
    finance_db.upsert_prediction_watch('fed-cut', 'Fed Cut Odds', 0.62, 'up', username=username)
    finance_db.set_market_brief_and_note('Market looks constructive.', 'Sentinel note.', username=username)


def test_render_chat_session_returns_payload_with_transcript_order():
    messages = _seed_chat()
    payload = render_chat_session('chat_export_1', 'Pantronux')

    assert payload.title == 'Export Session'
    assert payload.export_type == 'chat_session'
    assert payload.metadata['message_count'] == 3
    assert [row['id'] for row in payload.transcript] == [row['id'] for row in messages]
    assert payload.transcript[1]['content'].startswith('{')


def test_render_selected_messages_rejects_empty_message_ids():
    _seed_chat(chat_id='chat_export_2')
    try:
        render_selected_messages('chat_export_2', [], 'Pantronux')
    except HTTPException as exc:
        assert exc.status_code == 400
        return
    raise AssertionError('Expected HTTPException for empty message_ids')


def test_render_selected_messages_returns_only_requested_messages_in_order():
    messages = _seed_chat(chat_id='chat_export_3')
    selected_ids = [messages[2]['id'], messages[0]['id']]

    payload = render_selected_messages('chat_export_3', selected_ids, 'Pantronux')

    assert [row['id'] for row in payload.transcript] == sorted(selected_ids)
    assert payload.export_type == 'selected_messages'


def test_markdown_exporter_outputs_title_and_transcript():
    _seed_chat(chat_id='chat_export_md')
    payload = render_chat_session('chat_export_md', 'Pantronux')
    output = MarkdownExporter().export(payload).decode('utf-8')
    assert output.startswith('# Export Session')
    assert 'First user message' in output


def test_txt_exporter_outputs_plain_text_without_markdown_heading():
    _seed_chat(chat_id='chat_export_txt')
    payload = render_chat_session('chat_export_txt', 'Pantronux')
    output = TxtExporter().export(payload).decode('utf-8')
    assert not output.startswith('#')
    assert 'First user message' in output


def test_json_exporter_returns_valid_json():
    _seed_chat(chat_id='chat_export_json')
    payload = render_chat_session('chat_export_json', 'Pantronux')
    data = json.loads(JsonExporter().export(payload).decode('utf-8'))
    assert data['title'] == 'Export Session'
    assert len(data['transcript']) == 3


def test_pdf_exporter_returns_pdf_bytes():
    _seed_chat(chat_id='chat_export_pdf')
    payload = render_chat_session('chat_export_pdf', 'Pantronux')
    output = PdfExporter().export(payload)
    assert output.startswith(b'%PDF')
    assert len(output) > 100


def test_csv_exporter_returns_tabular_rows():
    _seed_chat(chat_id='chat_export_csv')
    payload = render_chat_session('chat_export_csv', 'Pantronux')
    output = CsvExporter().export(payload).decode('utf-8')
    assert output.startswith('id,timestamp,role')
    assert 'First user message' in output


def test_xlsx_exporter_returns_workbook_with_transcript_sheet():
    _seed_chat(chat_id='chat_export_xlsx')
    payload = render_chat_session('chat_export_xlsx', 'Pantronux')
    output = XlsxExporter().export(payload)
    workbook = load_workbook(filename=BytesIO(output))
    assert workbook.sheetnames == ['Metadata', 'Transcript']
    values = [workbook['Transcript'][f'F{row}'].value for row in range(2, 5)]
    assert 'First user message' in values
    assert 'Third user message' in values


def test_docx_exporter_returns_document_with_transcript_text():
    _seed_chat(chat_id='chat_export_docx')
    payload = render_chat_session('chat_export_docx', 'Pantronux')
    output = DocxExporter().export(payload)
    document = Document(BytesIO(output))
    text = '\n'.join(paragraph.text for paragraph in document.paragraphs)
    assert 'Export Session' in text
    assert 'First user message' in text


def test_validate_export_permission_rejects_foreign_chat_ownership():
    _seed_chat(chat_id='chat_export_owner', username='Pantronux')
    try:
        validate_export_permission('Faikhira', 'chat_export_owner')
    except HTTPException as exc:
        assert exc.status_code == 403
        return
    raise AssertionError('Expected HTTPException for foreign ownership')


def test_validate_export_permission_rejects_message_ids_from_other_chat():
    _seed_chat(chat_id='chat_export_a', username='Pantronux')
    other_messages = _seed_chat(chat_id='chat_export_b', username='Pantronux')
    bad_message_id = other_messages[0]['id']

    try:
        validate_export_permission('Pantronux', 'chat_export_a', [bad_message_id])
    except HTTPException as exc:
        assert exc.status_code == 403
        return
    raise AssertionError('Expected HTTPException for cross-chat message selection')


def test_sanitize_export_payload_removes_unsupported_keys():
    payload = ExportPayload(
        title='Unsafe',
        export_type='chat_session',
        username='Pantronux',
        source_chat_id='chat_export_unsafe',
        metadata={'exported_at': 'now', 'count': 2},
        transcript=[
            {
                'id': 1,
                'timestamp': '2026-01-01',
                'role': 'assistant',
                'persona': 'advisor',
                'role_label': 'Kuro (advisor)',
                'content': {'nested': True},
                'attachments': ['a.txt'],
                'debug': 'remove-me',
            }
        ],
    )

    cleaned = sanitize_export_payload(payload)
    assert 'debug' not in cleaned.transcript[0]
    assert isinstance(cleaned.transcript[0]['content'], str)
    assert cleaned.metadata['count'] == '2'


def test_render_intelligence_report_returns_payload():
    _seed_intelligence()
    payload = render_intelligence_report('Pantronux', briefing_date='2026-05-10')
    assert payload.export_type == 'intelligence_report'
    assert payload.metadata['briefing_date'] == '2026-05-10'
    assert any(table.title == 'stock_recommendations' for table in payload.tables)


def test_render_compliance_report_returns_payload():
    _seed_compliance()
    payload = render_compliance_report('Pantronux', standard='iso27001')
    assert payload.export_type == 'compliance_report'
    assert payload.metadata['standard'] == 'iso27001'
    assert any(table.title == 'evidence_matrix' for table in payload.tables)


def test_render_market_snapshot_returns_payload():
    _seed_market()
    payload = render_market_snapshot('Pantronux')
    assert payload.export_type == 'market_snapshot'
    assert payload.metadata['watched_symbol_count'] == 1
    assert any(table.title == 'market_hud_items' for table in payload.tables)
