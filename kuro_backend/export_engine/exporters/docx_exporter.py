from __future__ import annotations

from io import BytesIO

from docx import Document

from kuro_backend.export_engine.export_models import ExportPayload


class DocxExporter:
    format_name = "docx"
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    file_extension = ".docx"

    def export(self, payload: ExportPayload) -> bytes:
        document = Document()
        document.add_heading(payload.title, level=1)

        if payload.metadata:
            document.add_heading("Metadata", level=2)
            for key, value in payload.metadata.items():
                document.add_paragraph(f"{key}: {value}")

        document.add_heading("Transcript", level=2)
        for item in payload.transcript:
            role_label = item.get("role_label") or item.get("role", "unknown")
            timestamp = item.get("timestamp", "")
            document.add_heading(f"[{timestamp}] {role_label}", level=3)
            document.add_paragraph(item.get("content", ""))
            attachments = item.get("attachments") or []
            if attachments:
                document.add_paragraph("Attachments: " + ", ".join(attachments))

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
